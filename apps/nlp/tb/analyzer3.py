import nltk
import nltk.data
from nltk.corpus import brown
import string
from docx import *
import csv
import re
import pandas as pd
from sklearn import metrics
import os
import pickle
import scipy
import numpy
#import matplotlib.pyplot as plt
from scipy.cluster.hierarchy import dendrogram, linkage
import math
import unidecode
import json
from sklearn.metrics.pairwise import cosine_similarity
import operator
from sklearn.metrics import accuracy_score
from sklearn.metrics import f1_score

# path of the Word documents
# path = 'F:/vector_user/New corpus'
# filenames = os.listdir(path)

# path of authorship file
# author = 'author.txt'

# path of the output CSV document
# outpath_count = 'ratio_vector.csv'

# features besides pos, exmaple words, unnecessary words
otherfeature = ['TotalWords', 'AvgWordLength', 'AvgSentLength', 'ARI', 'boldwords', 'italicwords', 'bulletstyle',
                'Capwords', 'Text', 'Author']

examplewords = 'apps/nlp/tb/examplewords.txt'
unnecessary = 'apps/nlp/tb/unnecessary.txt'
author = 'apps/nlp/tb/author.txt'
zscore_vector = 'apps/nlp/tb/zscore_vector.csv'


examplewords_f = open(examplewords, 'r')
examplewords_text = examplewords_f.read()
examplewords_list = examplewords_text.split('\n')

# read the words in the word document
def getWord(document):
    doc = []
    for paragraph in document.paragraphs:
        if len(paragraph.text.split()) > 10:
            doc.append([unidecode.unidecode(paragraph.text)])
            doc[len(doc) - 1].append(paragraph.style.name)
            doc[len(doc) - 1].append(0)
            doc[len(doc) - 1].append(0)
            for run in paragraph.runs:
                if run.bold:
                    doc[len(doc) - 1][2] += len(nltk.word_tokenize(run.text))
                if run.italic:
                    doc[len(doc) - 1][3] += len(nltk.word_tokenize(run.text))
    return doc


def getWord1(document):
    doc1 = []
    for paragraph in document.paragraphs:
        if len(paragraph.text.split()) > 10:
            doc1.append(paragraph.text)
    return doc1


# Transform the text into POS tags
def pos(text_tagged):
    # lower, tokenize, and POS the text; may need some recoding of special words
    # pattern = re.compile('.*[^a-z0-9].*')
    text_transf = ''
    p = re.compile('\w')
    for pair in text_tagged:
        if p.match(pair[1][0]):
            text_transf += pair[1] + ' '
    return text_transf


# Transform the text into POS vector
def pos_uni_features(text, pos_tags):
    features = {}
    unifdist = nltk.FreqDist(nltk.word_tokenize(text))
    for tag in pos_tags:
        features['%s' % tag] = unifdist[tag]
    return features


# Transform the text given into bigram vector
# def pos_bi_features(text,bipos):
#	features = {}
#	bifdist = nltk.FreqDist(nltk.bigrams(nltk.word_tokenize(text)))
#	for word in bipos:
#		features['%s' % word] = bifdist[word.split()[0],word.split()[1]]
#	return features

# Transform the text into word vector
def word_features(text, wordlist):
    features = {}
    unifdist = nltk.FreqDist(nltk.word_tokenize(text.lower()))
    bifdist = nltk.FreqDist(nltk.bigrams(nltk.word_tokenize(text.lower())))
    for word in wordlist:
        if len(nltk.word_tokenize(word)) == 1:
            features['%s' % word] = unifdist[word]
        if len(nltk.word_tokenize(word)) == 2:
            features['%s' % word] = bifdist[nltk.word_tokenize(word)[0], nltk.word_tokenize(word)[1]]
        if len(nltk.word_tokenize(word)) > 2:
            features['%s' % word] = text.lower().count(word)
    return features


# Bieber features
def bieber(text_tagged):
	features = {}
	thatComplement = 0
	publicVerb = 0
	privateVerb = 0
	firstPersonPronouns = 0
	possibilityModal = 0
	downtoners = 0
	AgentlessPassives = 0
	for i in range(len(text_tagged)):
		if text_tagged[i][0].lower() == 'that' and text_tagged[i][1] == 'IN':
			thatComplement+=1
		if text_tagged[i][0].lower() in ['affirm', 'announce', 'boast', 'confirm', 'decalre']:
			publicVerb+=1
		if text_tagged[i][0].lower() in ['believe', 'know', 'realize', 'understand']:
			privateVerb+=1
		if text_tagged[i][0].lower() in ['i','me','my','mine','myself','we','our','ours','us','ourselves']:
			firstPersonPronouns+=1
		if text_tagged[i][0].lower() in ['may','might','could','must']:
			possibilityModal+=1
		if text_tagged[i][0].lower() in ['hardly','slightly','barely','just','somewhat']:
			downtoners+=1
		if text_tagged[i][0].lower().endswith('ish') and len(text_tagged[i][0].lower())>5:
			downtoners+=1
	for i in range(len(text_tagged)-1):
		if text_tagged[i][0].lower() in ['have','has','had'] and text_tagged[i+1][0].lower() == 'to':
			possibilityModal+=1
		if text_tagged[i][0].lower() + ' ' + text_tagged[i+1][0].lower() in ['a bit','a little','only just''kind of','sort of','little bit','tiny bit']:
			downtoners+=1

	features['thatComplement'] = thatComplement
	features['publicVerb'] = publicVerb
	features['privateVerb'] = privateVerb
	features['firstPersonPronouns'] = firstPersonPronouns
	features['possibilityModal'] = possibilityModal
	features['downtoners'] = downtoners
	return features


# Compute average word length per paragraph
def avg_word_length(text):
    totalchar = 0
    text = re.sub(r'[^\w\s]', '', text)
    wordcount = len(text.split())
    for word in text.split():
        totalchar += len(word)
    return float(totalchar) / float(wordcount)


# Ccompute average sentence length per paragraph
def ave_sent_length(text):
    sent_detector = nltk.data.load('tokenizers/punkt/english.pickle')
    sent_list = sent_detector.tokenize(text.strip())
    sentcount = len(sent_list)
    sentwords = len(text.split())
    return float(sentwords) / float(sentcount)


# Compute ARI readibility score
def ari_score(text):
    return 4.71 * avg_word_length(text) + (0.5 * ave_sent_length(text)) - 21.43


# Output a CSV document of the count vector
def cvswrite_count(featuresets1, featuresets2, featuresets3, text, outpath):
    fp = open(outpath, "wb")
    writer = csv.writer(fp)
    # pattern = re.compile('.*[^a-z].*')
    # write the feature names
    featurenames = sorted(featuresets1[0].keys()) + sorted(featuresets2[0].keys()) + sorted(
        featuresets3[0].keys()) + otherfeature
    writer.writerow(featurenames)
    # write values for each paragraph as a row
    for i in range(0, len(featuresets1)):
        featureline = []
        # write values of POS features
        for key in sorted(featuresets1[0].keys()):
            featureline.append(str(featuresets1[i][key]))
        # write values of example words features
        for key in sorted(featuresets2[0].keys()):
            featureline.append(str(featuresets2[i][key]))
        # write values of bieber features
        for key in sorted(featuresets3[0].keys()):
            featureline.append(str(featuresets3[i][key]))
            # for key in sorted(featuresets3[0].keys()):
            # featureline.append(str(featuresets3[i][key]))
        # write word count for each paragraph
        featureline.append(len(text[i][0].split()))
        # write average word length
        featureline.append(avg_word_length(text[i][0]))
        # write average sentence length
        featureline.append(ave_sent_length(text[i][0]))
        # write ARI
        featureline.append(ari_score(text[i][0]))
        # write style features
        featureline.append(text[i][2])
        featureline.append(text[i][3])
        featureline.append('List' in text[i][1])
        featureline.append(sum(1 for word in nltk.word_tokenize(text[i][0]) if word.isupper() and len(word) > 1) > 0)
        # write content of the paragraph
        featureline.append(text[i][0].encode('utf-8'))
        # write author
        #featureline.append(str(author_true[i]))
        writer.writerow(featureline)
    fp.close()

def analyze3(infile):

    # all the POS tags
    postags = ['CC', 'CD', 'DT', 'EX', 'FW', 'IN', 'JJ', 'JJR', 'JJS', 'LS', 'MD', 'NN', 'NNS', 'NNP', 'NNPS', 'PDT',
               'POS', 'PRP', 'PRP$', 'RB', 'RBR', 'RBS', 'RP', 'SYM', 'TO', 'UH', 'VB', 'VBD', 'VBG', 'VBN', 'VBP',
               'VBZ', 'WDT', 'WP', 'WP$', 'WRB']

    # bipos_file = open('bipos.txt','r')
    # bipos_text = bipos_file.read()
    # bipos = bipos_text.split('\n')

    # read the word list
    # word_file = open('wordlist.txt','r')
    word_file = 'apps/nlp/tb/wordlist.txt'
    word_filef = open(word_file, 'r')
    word_text = word_filef.read()
    wordlist = word_text.split('\n')
    wordlist = list(set(wordlist))
    print len(wordlist)

    # generate brown words freq dist

    brown_pos_fdist = json.load(open("apps/nlp/tb/brown_pos_fdist.txt"))
    brown_word_fdist = json.load(open("apps/nlp/tb/brown_word_fdist.txt"))

    # start processing each document

    outpath_count = 'apps/nlp/tb/' + '/' + 'count_vector.csv'

    # read the MS word document into a string list
    document = Document(infile)
    doc = getWord(document)
    doc1 = getWord1(document)

    # Get POS features
    text_tagged = [nltk.pos_tag(nltk.word_tokenize(line)) for line, style, bold, italic in doc]
    postext = [(pos(text)) for text in text_tagged]
    featuresets1 = [(pos_uni_features(line, postags)) for line in postext]
    # featuresets2 = [(pos_bi_features(line,bipos)) for line in postext]

    # Get words feature
    featuresets2 = [(word_features(line, wordlist)) for line, style, bold, italic in doc]
    featuresets3 = [(bieber(line)) for line in text_tagged]
    # if all the featuresets have equal length with length of the document, then write the csv file
    if len(featuresets1) == len(featuresets2):
         cvswrite_count(featuresets1, featuresets2, featuresets3, doc, outpath_count)

    # drop the feature that occurs only once in a document
    df_ratio = pd.read_csv(outpath_count, sep=',')
    cols = list(df_ratio)
    for col in cols[0:len(cols) - 2]:
        if df_ratio[col].sum() < 2:
            df_ratio = df_ratio.drop(col, 1)

    # ratio calculation
    cols = list(df_ratio)
    for col in cols:
        if col in brown_pos_fdist.keys():
            df_ratio[col] = (df_ratio[col] * 100 / df_ratio['TotalWords']) / (math.log(brown_pos_fdist[col] + 1) + 0.5)
        if col in brown_word_fdist.keys():
            df_ratio[col] = (df_ratio[col] * 100 / df_ratio['TotalWords']) / (math.log(brown_word_fdist[col] + 1) + 0.5)
        if col == 'boldwords' or col == 'italicwords':
            df_ratio[col] = df_ratio[col] = df_ratio[col] * 100 / df_ratio['TotalWords']

    df_ratio.index += 1
    # df_ratio.to_csv(path + '/' + filename + '/' + 'ratio_vector.csv')

    # z score
    for col in cols[0:len(cols) - 2]:
        if df_ratio[col].std(ddof=0) != 0:
            df_ratio[col] = (df_ratio[col] - df_ratio[col].mean()) / df_ratio[col].std(ddof=0)
        else:
            df_ratio[col] = 0

    # df_ratio.to_csv(path + '/' + filename + '/' + 'zscore_vector.csv')

    # to compute difference between two adjacent paragraphs
    df_sub = df_ratio.ix[:, 0:len(df_ratio.columns) - 2]
    df_sub['AuthorChange'] = df_ratio['Author']
    vector_diff = df_sub.values
    df_new = pd.DataFrame(columns=list(df_sub))

    for i in range(1, len(vector_diff)):
        df_new.loc[i - 1] = vector_diff[i - 1] - vector_diff[i]

    # add paragraph index of two adjacent paragraphs
    df_new['PreviousParagraph'] = [i for i in range(1, len(df_ratio))]
    df_new['NextsParagraph'] = [i for i in range(2, len(df_ratio) + 1)]

    # compute the cosine similarity
    vector_feature = df_ratio.ix[:, 0:len(df_ratio.columns) - 2].values
    cos_matrix = cosine_similarity(vector_feature)
    cos = []
    for i in range(len(df_ratio) - 1):
        cos.append(cos_matrix[i][i + 1])

    df_new['CosineSimilarity'] = cos

    # find features whose absolute difference between 2 adjacent paragraphs greater than 2
    # and there is a cosine similarity smaller than 0
    bigdiff_feature = []
    feature_count = len(df_new.ix[:, 0:len(df_new.columns) - 4].values[0])
    words = []
    print len(df_new), len(text_tagged)
    for i in range(0, len(df_new)):
        # print 'new'
        dic = {}
        bigdiffwords = []
        for j in range(feature_count):
            if abs(df_new.ix[:, 0:len(df_new.columns) - 4].values[i][j]) >= 2 and df_new['CosineSimilarity'][i] < 0:
                dic[list(df_new)[j]] = [df_new.ix[:, 0:len(df_new.columns) - 4].values[i][j]]
                if list(df_new)[j] in postags:
                    for m in range(len(text_tagged[i])):
                        if text_tagged[i][m][1] == list(df_new)[j]:
                            dic[list(df_new)[j]].append(text_tagged[i][m][0])
                    for n in range(len(text_tagged[i + 1])):
                        if text_tagged[i + 1][n][1] == list(df_new)[j]:
                            dic[list(df_new)[j]].append(text_tagged[i + 1][n][0])
                else:
                    dic[list(df_new)[j]].append(list(df_new)[j])
        words.append(dic)
        templist = sorted(dic.items(), key=operator.itemgetter(1), reverse=True)
        for pair in templist:
            bigdiffwords.append(pair[0])
        # print pair[0],pair[1]
        bigdiff_feature.append(bigdiffwords)

    df_new['BigDiffFeatures'] = ['|'.join(diff) for diff in bigdiff_feature]
    print 'bigdiff_feature', bigdiff_feature

    # output the df_new as a csv file
    # df_new.to_csv(path + '/' + filename + '/' + 'similarity.csv')

    # the words to highlight as big difference
    print 'words', words

    para = 0
    for word_dict in words:
        dkeys = word_dict.keys()
        # print "paragraph %s" % para
        # print doc[para]
        for dkey in dkeys:
            for hiword in word_dict[dkey]:
                if hiword == '.':
                    continue
                classname = dkey + '_' + '%s' % para
                doc1[para] = re.sub(r'\b%s\b' % hiword, '<span class ="%s">%s</span>' % (classname, hiword), doc1[para],
                                    flags=re.UNICODE)
                a = para + 1
                classname = dkey + '_' + '%s' % a
                doc1[para + 1] = re.sub(r'\b%s\b' % hiword, '<span class ="%s">%s</span>' % (classname, hiword),
                                        doc1[para + 1], flags=re.UNICODE)
        para = para + 1


    #create alert boxes

    alerts = []
    para = 0
    for word_dict in words:
        s = []
        strm =[]
        strl =[]
        content = ''
        dkeys = word_dict.keys()
        for dkey in dkeys:
            hiword = word_dict[dkey]
            if dkey == 'TotalWords':
                if hiword[0]>0:
                    str = 'The paragraph above has many more words than the one below. '
                else:
                    str = 'The paragraph above has many fewer words than the one below. '
                s.append(str)
            elif dkey == 'AvgSentLength':
                if hiword[0] > 0:
                    str = 'The paragraph above has longer sentences than the one below. '
                else:
                    str = 'The paragraph above has shorter sentences than the one below. '
                s.append(str)
            elif dkey == 'AvgWordLength':
                if hiword[0] > 0:
                    str = 'The paragraph above has longer words than the one below. '
                else:
                    str = 'The paragraph above has shorter words than the one below. '
                s.append(str)
            else:
                classname = dkey + '_' + '%s' % para
                featureclass = classname + 'f'
                if hiword[0] > 0:
                    strmore = '<span onMouseOver="setfeaturecolor(\'%s\')" onmouseout="onMouseOut(\'%s\')" id ="%s">%s</span> ' % (
                    classname, classname, featureclass, dkey)
                    strm.append(strmore)
                else:
                    strless = '<span onMouseOver="setfeaturecolor(\'%s\')" onmouseout="onMouseOut(\'%s\')" id ="%s">%s</span> ' % (
                        classname, classname, featureclass, dkey)
                    strl.append(strless)
        strmore = ', '.join(strm)
        strless = ', '.join(strl)
        print 'strmore', strmore
        print 'strless', strless
        str = ''
        if strmore:
            str = 'The paragraph above has many more %s than the one below.' % strmore
        if strless:
            str += ' The paragraph above has many less %s than the one below. ' % strless
        s.append(str)
        s = ' '.join(s)
        if s:
            content = '<div class="alert"> <span class="closebtn" onclick="this.parentElement.style.display=\'none\';">&times;</span>%s </div>' % s
        alerts.append(content)
        para += 1



    result = []
    for i in range(len(words)):
        result.append(doc1[i])
        if (words[i] != {}):
            result.append(alerts[i])
    result.append(doc1[i + 1])

    return result

    # result_file.close()


